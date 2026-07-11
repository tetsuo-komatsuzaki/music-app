"""md_to_docx.py — 簡易 Markdown → Word(.docx) 変換（見出し/段落/箇条書き/表 対応・日本語フォント設定）

使い方: python md_to_docx.py <input.md> <output.docx>
"""
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

JP_FONT = "Yu Gothic"


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), JP_FONT)


def clean_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)  # bold
    s = re.sub(r"`(.+?)`", r"\1", s)         # code
    s = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", s)  # links
    return s.strip()


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def parse_cells(line: str):
    return [clean_inline(c) for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    header = parse_cells(rows[0])
    data = [parse_cells(r) for r in rows[2:]]  # rows[1] = separator
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for drow in data:
        cells = t.add_row().cells
        for i in range(len(header)):
            cells[i].text = drow[i] if i < len(drow) else ""
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def main(in_path, out_path):
    with open(in_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    set_default_font(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # 表ブロック
        if is_table_row(line) and i + 1 < len(lines) and set(lines[i + 1].strip()) <= set("|-: "):
            block = []
            while i < len(lines) and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            doc.add_paragraph("")
            continue

        # 見出し
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(clean_inline(m.group(2)), level=min(level, 4))
            i += 1
            continue

        # 箇条書き
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            doc.add_paragraph(clean_inline(m.group(1)), style="List Bullet")
            i += 1
            continue

        # 番号付き
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            doc.add_paragraph(clean_inline(m.group(1)), style="List Number")
            i += 1
            continue

        doc.add_paragraph(clean_inline(stripped))
        i += 1

    doc.save(out_path)
    print(f"saved: {out_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
